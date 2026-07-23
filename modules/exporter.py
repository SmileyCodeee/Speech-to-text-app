"""
exporter.py
------------
Saves generated notes to shareable formats: .txt, .docx, .pdf, .md

Uses:
  - python-docx  -> https://python-docx.readthedocs.io/
  - reportlab    -> https://www.reportlab.com/docs/reportlab-userguide.pdf
"""

import io
from typing import List, Optional

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem


def to_txt(title: str, paragraphs: List[str], summary: str = "", keywords: Optional[List[str]] = None) -> bytes:
    lines = [title, "=" * len(title), ""]
    if summary:
        lines += ["SUMMARY:", summary, ""]
    if keywords:
        lines += ["KEYWORDS: " + ", ".join(keywords), ""]
    lines += ["NOTES:", ""]
    lines += paragraphs
    return "\n\n".join(lines).encode("utf-8")


def to_markdown(title: str, paragraphs: List[str], summary: str = "", keywords: Optional[List[str]] = None) -> bytes:
    lines = [f"# {title}", ""]
    if summary:
        lines += ["## Summary", summary, ""]
    if keywords:
        lines += ["## Keywords", ", ".join(f"`{k}`" for k in keywords), ""]
    lines += ["## Notes", ""]
    for p in paragraphs:
        lines.append(f"- {p}")
    return "\n".join(lines).encode("utf-8")


def to_docx(title: str, paragraphs: List[str], summary: str = "", keywords: Optional[List[str]] = None) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)

    if summary:
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(summary)

    if keywords:
        doc.add_heading("Keywords", level=2)
        doc.add_paragraph(", ".join(keywords))

    doc.add_heading("Notes", level=2)
    for p in paragraphs:
        para = doc.add_paragraph(p, style="List Bullet")
        para.paragraph_format.space_after = Pt(6)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(title: str, paragraphs: List[str], summary: str = "", keywords: Optional[List[str]] = None) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    h1 = styles["Title"]
    h2 = styles["Heading2"]
    body = ParagraphStyle("Body", parent=styles["BodyText"], spaceAfter=8, leading=15)

    story = [Paragraph(title, h1), Spacer(1, 12)]

    if summary:
        story += [Paragraph("Summary", h2), Paragraph(summary, body), Spacer(1, 10)]

    if keywords:
        story += [Paragraph("Keywords", h2), Paragraph(", ".join(keywords), body), Spacer(1, 10)]

    story.append(Paragraph("Notes", h2))
    items = [ListItem(Paragraph(p, body)) for p in paragraphs]
    story.append(ListFlowable(items, bulletType="bullet"))

    doc.build(story)
    return buf.getvalue()


EXPORTERS = {
    "txt": to_txt,
    "md": to_markdown,
    "docx": to_docx,
    "pdf": to_pdf,
}


def export(fmt: str, title: str, paragraphs: List[str], summary: str = "", keywords: Optional[List[str]] = None) -> bytes:
    fmt = fmt.lower().lstrip(".")
    if fmt not in EXPORTERS:
        raise ValueError(f"Unsupported format: {fmt}. Choose from {list(EXPORTERS)}")
    return EXPORTERS[fmt](title, paragraphs, summary, keywords)